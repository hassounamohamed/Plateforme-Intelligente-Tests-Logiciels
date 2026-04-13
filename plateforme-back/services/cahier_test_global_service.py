"""
Service du Cahier de Tests Global.

Fonctionnalités :
  1. Démarrer la génération (non-bloquant) → retourne AIGeneration immédiatement
  2. Exécuter la génération en arrière-plan (BackgroundTask) avec logs dans ai_logs
  3. Valider le cahier
  4. Mettre à jour un cas de test
  5. Exporter en PDF, Word, Excel
"""
from __future__ import annotations

import io
import json
import os
import re
import time
import unicodedata
from datetime import datetime
from typing import List, Optional

from fastapi import HTTPException, status
from sqlalchemy.orm import Session, joinedload

from core.config import AI_API_KEY, AI_MODEL, AI_API_URL
from core.rbac.constants import ROLE_DEVELOPPEUR, ROLE_TESTEUR_QA
from models.notification import TypeNotification
from models.ai_generation import AIGeneration
from models.rapports import RapportQA, IndicateurQualite, RecommandationQualite
from models.scrum import Projet, Sprint, UserStory, Module, Epic
from models.cahier_test_global import CahierTestGlobal, CasTest
from models.user import Utilisateur
from repositories.ai_generation_repository import AIGenerationRepository
from repositories.cahier_test_global_repository import CahierTestGlobalRepository
from services.notification_service import NotificationService
from schemas.cahier_test_global import (
    AICahierResponse,
    CreateCasTestRequest,
    UpdateCasTestRequest,
)

# ─── Prompt système ───────────────────────────────────────────────────────────

SYSTEM_PROMPT = """\
Tu es un expert en qualité logicielle et en tests. Ton rôle est de générer \
un cahier de tests complet et structuré pour un projet Scrum à partir \
de la liste de ses sprints et user stories.

Règles strictes :
1. Pour chaque user story, génère entre 2 et 5 cas de tests couvrant \
   les scénarios nominaux, alternatifs et d'erreur.
2. Chaque cas de test doit contenir :
    - user_story_id   : identifiant numérique exact de la user story source (obligatoire)
   - sprint          : nom du sprint (ex: "Sprint 1")
    - module          : nom du module lié à la user story (optionnel)
    - sous_module     : sous-composant ou fonctionnalité précise (optionnel)
   - test_ref        : référence unique au format TC-XXX (ex: TC-001)
   - test_case       : titre court et descriptif du cas de test
   - test_purpose    : objectif du test en une phrase
   - type_utilisateur: type d'utilisateur impliqué (ex: Admin, Utilisateur, Testeur)
   - scenario_test   : étapes numérotées du scénario (ex: "1. Ouvrir... 2. Saisir...")
   - resultat_attendu: résultat observable attendu après exécution
    - execution_time_seconds: durée estimée d'exécution en secondes (entier >= 0)
   - type_test       : "Manuel" ou "Automatisé"
3. Les test_ref doivent être séquentiels et uniques (TC-001, TC-002, …).
4. Le user_story_id doit obligatoirement correspondre à une user story fournie dans le contexte.
5. Retourne UNIQUEMENT un objet JSON valide, aucun texte avant ou après.

Structure JSON attendue :
{
  "cas_tests": [
    {
            "user_story_id": 14,
      "sprint": "Sprint 1",
      "module": "Authentification",
      "sous_module": "Connexion",
      "test_ref": "TC-001",
      "test_case": "Connexion avec identifiants valides",
      "test_purpose": "Vérifier que l'utilisateur peut se connecter avec des identifiants corrects",
      "type_utilisateur": "Utilisateur",
      "scenario_test": "1. Ouvrir la page de connexion\\n2. Saisir l'email valide\\n3. Saisir le mot de passe correct\\n4. Cliquer sur Connexion",
      "resultat_attendu": "L'utilisateur est redirigé vers le tableau de bord",
            "execution_time_seconds": 30,
      "type_test": "Manuel"
    }
  ]
}
"""

USER_PROMPT_TEMPLATE = """\
Voici les informations du projet Scrum :

Projet : {projet_nom}
Description : {projet_description}

Sprints et User Stories :
{sprints_content}

Génère le cahier de tests complet en JSON selon le schéma demandé.
"""

BUG_SUGGESTION_SYSTEM_PROMPT = """\
Tu es un expert QA et triage de bugs.
À partir du contexte d'un cas de test en échec ou bloqué, propose:
1) bug_titre_correction: un titre court, actionnable, orienté correction.
2) bug_nom_tache: un nom de tâche concis de type backlog/issue.

Contraintes:
- Réponds UNIQUEMENT en JSON valide.
- Format strict:
{
    "bug_titre_correction": "...",
    "bug_nom_tache": "..."
}
- Longueur max 120 caractères par champ.
"""

RAPPORT_QA_SYSTEM_PROMPT = """\
Tu es un lead QA.
Tu reçois les statistiques d'un cahier de tests global.
Ta mission est de générer un rapport QA concis et actionnable.

Retourne UNIQUEMENT un JSON valide, sans texte autour, au format strict :
{
    "statut": "brouillon|valide",
    "recommandations": "texte multi-lignes",
    "tendance": "amelioration|stable|degradation",
    "indice_qualite": 0.0,
    "nombre_anomalies_critiques": 0
}
"""


class CahierTestGlobalService:

    def __init__(self, db: Session):
        self.db      = db
        self.repo    = CahierTestGlobalRepository(db)
        self.ai_repo = AIGenerationRepository(db)
        self.notification_service = NotificationService(db)
        self.api_key_service = None

    def _get_api_key_service(self):
        """Lazy load APIKeyService to avoid circular imports."""
        if self.api_key_service is None:
            from services.api_key_service import APIKeyService
            self.api_key_service = APIKeyService(self.db)
        return self.api_key_service

    def _get_api_key_for_request(self, user_id: Optional[int]) -> str:
        """
        Resolve API key for AI call.

        Priority:
        1) User custom key (if enabled)
        2) Shared platform key from environment
        """
        if user_id:
            api_key_svc = self._get_api_key_service()
            custom_key = api_key_svc.get_api_key_for_user(user_id)
            if custom_key:
                return custom_key

        if not AI_API_KEY:
            raise ValueError("Clé API IA manquante. Ajoutez votre clé API dans le profil ou définissez AI_API_KEY dans .env")

        return AI_API_KEY

    @staticmethod
    def _parse_execution_time_seconds(value: object) -> Optional[int]:
        if value is None:
            return None
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return None
        return max(0, parsed)

    def _resolve_user_story_id(
        self,
        projet_id: int,
        user_story_id: Optional[int],
        module: Optional[str] = None,
        sous_module: Optional[str] = None,
        test_case: Optional[str] = None,
    ) -> int:
        us_query = (
            self.db.query(UserStory)
            .join(Epic, UserStory.epic_id == Epic.id)
            .join(Module, Epic.module_id == Module.id)
            .filter(Module.projet_id == projet_id)
        )

        if user_story_id is not None:
            us = us_query.filter(UserStory.id == user_story_id).first()
            if not us:
                raise HTTPException(
                    status_code=422,
                    detail="user_story_id invalide: la user story n'appartient pas au projet.",
                )
            return us.id

        user_stories = us_query.all()
        if not user_stories:
            raise HTTPException(
                status_code=422,
                detail="Aucune user story disponible dans ce projet pour lier le cas de test.",
            )

        module_l = (module or "").strip().lower()
        sous_module_l = (sous_module or "").strip().lower()
        test_case_l = (test_case or "").strip().lower()
        for us in user_stories:
            title_l = (us.titre or "").lower()
            desc_l = (us.description or "").lower()
            module_name_l = ((us.epic.module.nom if us.epic and us.epic.module else "") or "").lower()

            if module_l and module_l in module_name_l:
                return us.id
            if sous_module_l and sous_module_l in title_l:
                return us.id
            if test_case_l and (test_case_l in title_l or (title_l and title_l in test_case_l) or test_case_l in desc_l):
                return us.id

        return user_stories[0].id

    @staticmethod
    def _attach_user_story_display(cas: CasTest) -> CasTest:
        user_story = getattr(cas, "user_story", None)
        cas.user_story_reference = user_story.reference if user_story else None
        cas.user_story_titre = user_story.titre if user_story else None
        return cas

    def _attach_user_story_display_many(self, cases: List[CasTest]) -> List[CasTest]:
        return [self._attach_user_story_display(cas) for cas in cases]

    # ── Points d'entrée publics ──────────────────────────────────────────

    def demarrer_generation(
        self,
        projet_id: int,
        user_id: int,
        version: str = "1.0.0",
        mode_generation: str = "ai",
    ) -> AIGeneration | CahierTestGlobal:
        """
        Crée le job AIGeneration + initialise le CahierTestGlobal.
        Retourne immédiatement le job (status=pending).
        L'appelant doit lancer `executer_generation` en BackgroundTask.
        """
        projet = self.db.query(Projet).filter(Projet.id == projet_id).first()
        if not projet:
            raise HTTPException(status_code=404, detail="Projet introuvable.")

        # Mode manuel : initialise un cahier vide, sans job IA.
        if mode_generation == "manuelle":
            cahier = self.repo.get_by_projet(projet_id)
            if cahier:
                self.repo.delete_cas_tests(cahier.id)
                cahier.version = version
                cahier.statut = "brouillon"
                cahier.date_generation = datetime.utcnow()
                cahier.generated_by_id = user_id
                cahier.ai_generation_id = None
                cahier.nombre_total = 0
                cahier.nombre_reussi = 0
                cahier.nombre_echoue = 0
                cahier.nombre_bloque = 0
                self.db.commit()
                self.db.refresh(cahier)
            else:
                cahier = CahierTestGlobal(
                    projet_id=projet_id,
                    generated_by_id=user_id,
                    version=version,
                    statut="brouillon",
                    ai_generation_id=None,
                    nombre_total=0,
                    nombre_reussi=0,
                    nombre_echoue=0,
                    nombre_bloque=0,
                )
                self.db.add(cahier)
                self.db.commit()
                self.db.refresh(cahier)
            return cahier

        # Créer le job de génération IA (type=generate_tests)
        gen = self.ai_repo.create_generation(projet_id, user_id, "generate_tests")

        # Créer ou réinitialiser le cahier
        cahier = self.repo.get_by_projet(projet_id)
        if cahier:
            self.repo.delete_cas_tests(cahier.id)
            cahier.version          = version
            cahier.statut           = "generating"
            cahier.date_generation  = datetime.utcnow()
            cahier.generated_by_id  = user_id
            cahier.ai_generation_id = gen.id
            cahier.nombre_total     = 0
            cahier.nombre_reussi    = 0
            cahier.nombre_echoue    = 0
            cahier.nombre_bloque    = 0
            self.db.commit()
            self.db.refresh(cahier)
        else:
            cahier = self.repo.create_cahier(projet_id, user_id, version, ai_generation_id=gen.id)

        return gen

    def create_cas_test(
        self,
        cahier_id: int,
        projet_id: int,
        data: CreateCasTestRequest,
    ) -> CasTest:
        """Créer un cas de test manuel dans un cahier existant."""
        self._verifier_appartenance(cahier_id, projet_id)

        last_cas = (
            self.db.query(CasTest)
            .filter(CasTest.cahier_id == cahier_id)
            .order_by(CasTest.ordre.desc())
            .first()
        )
        next_order = (last_cas.ordre + 1) if last_cas else 1
        resolved_user_story_id = self._resolve_user_story_id(
            projet_id=projet_id,
            user_story_id=data.user_story_id,
            module=data.module,
            sous_module=data.sous_module,
            test_case=data.test_case,
        )

        cas = self.repo.add_cas_test(
            cahier_id=cahier_id,
            user_story_id=resolved_user_story_id,
            sprint=data.sprint or "",
            module=data.module or "",
            sous_module=data.sous_module or "",
            test_ref=f"TC-{next_order:03d}",
            test_case=data.test_case,
            test_purpose=data.test_purpose or "",
            type_utilisateur=data.type_utilisateur or "",
            scenario_test=data.scenario_test or "",
            resultat_attendu=data.resultat_attendu or "",
            execution_time_seconds=data.execution_time_seconds,
            type_test=data.type_test,
            ordre=next_order,
        )

        if data.commentaire is not None:
            cas.commentaire = data.commentaire
            self.db.commit()
            self.db.refresh(cas)

        if data.execution_time_seconds is not None:
            cas.execution_time_seconds = data.execution_time_seconds
            self.db.commit()
            self.db.refresh(cas)

        # Toute création manuelle d'un cas incrémente la version mineure du cahier.
        cahier = self.repo.get_by_id(cahier_id)
        if cahier:
            cahier.version = self._increment_cahier_minor_version(cahier.version)

        self.repo.recalculer_stats(cahier_id)
        return self._attach_user_story_display(cas)

    def _increment_cahier_minor_version(self, current_version: Optional[str]) -> str:
        """
        Incrémente la version mineure du cahier : X.Y.Z -> X.(Y+1).0
        Exemple: 1.0.0 -> 1.1.0
        """
        default_major, default_minor, default_patch = 1, 0, 0
        if not current_version:
            major, minor, patch = default_major, default_minor, default_patch
        else:
            match = re.fullmatch(r"(\d+)\.(\d+)\.(\d+)", current_version.strip())
            if not match:
                major, minor, patch = default_major, default_minor, default_patch
            else:
                major, minor, patch = map(int, match.groups())

        return f"{major}.{minor + 1}.0"

    def executer_generation(self, generation_id: int) -> None:
        """
        Logique principale exécutée en arrière-plan.
        Toutes les étapes sont loguées dans ai_logs.
        """
        try:
            self._run(generation_id)
        except Exception as exc:
            # IMPORTANT: remettre la session en état avant toute écriture de statut/log
            self.db.rollback()
            try:
                self.ai_repo.update_status(generation_id, "failed", 0)
                self.ai_repo.add_log(generation_id, "error", str(exc), 0)
                # Mettre le cahier en erreur
                gen = self.ai_repo.get_by_id(generation_id)
                if gen:
                    cahier = self.repo.get_by_projet(gen.projet_id)
                    if cahier and cahier.ai_generation_id == generation_id:
                        cahier.statut = "failed"
                        self.db.commit()
            except Exception:
                # Ne pas relancer ici pour éviter de perdre la cause d'origine.
                self.db.rollback()

    def get_generation(self, generation_id: int, projet_id: int) -> AIGeneration:
        gen = self.ai_repo.get_detail(generation_id)
        if not gen or gen.projet_id != projet_id or gen.type != "generate_tests":
            raise HTTPException(status_code=404, detail="Génération introuvable.")
        return gen

    def list_generations(self, projet_id: int) -> List[AIGeneration]:
        return (
            self.db.query(AIGeneration)
            .filter(
                AIGeneration.projet_id == projet_id,
                AIGeneration.type == "generate_tests",
            )
            .order_by(AIGeneration.created_at.desc())
            .all()
        )

    def valider_cahier(self, cahier_id: int, projet_id: int, version: Optional[str]) -> CahierTestGlobal:
        self._verifier_appartenance(cahier_id, projet_id)
        return self.repo.valider(cahier_id, version)

    def update_cas_test(
        self,
        cahier_id: int,
        cas_id: int,
        projet_id: int,
        data: UpdateCasTestRequest,
        changed_by_id: Optional[int] = None,
    ) -> CasTest:
        self._verifier_appartenance(cahier_id, projet_id)
        cas = self.repo.get_cas_test(cas_id, cahier_id)
        if not cas:
            raise HTTPException(status_code=404, detail="Cas de test introuvable.")
        before = {
            "statut_test": cas.statut_test,
            "type_test": cas.type_test,
            "commentaire": cas.commentaire,
            "bug_titre_correction": cas.bug_titre_correction,
            "bug_nom_tache": cas.bug_nom_tache,
        }
        payload = data.model_dump(exclude_none=True)

        statut_cible = payload.get("statut_test", cas.statut_test)
        bug_titre = payload.get("bug_titre_correction", cas.bug_titre_correction)
        bug_tache = payload.get("bug_nom_tache", cas.bug_nom_tache)

        if statut_cible in {"Échoué", "Bloqué"}:
            if not (bug_titre and bug_titre.strip()) or not (bug_tache and bug_tache.strip()):
                suggested_title, suggested_task = self._generer_bug_fields_ia(cas, None)
                payload.setdefault("bug_titre_correction", suggested_title)
                payload.setdefault("bug_nom_tache", suggested_task)
                bug_titre = payload.get("bug_titre_correction")
                bug_tache = payload.get("bug_nom_tache")
            if not (bug_titre and bug_titre.strip()):
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Le titre de correction est obligatoire pour un test échoué ou bloqué.",
                )
            if not (bug_tache and bug_tache.strip()):
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Le nom de tâche est obligatoire pour un test échoué ou bloqué.",
                )

        updated = self.repo.update_cas_test(cas_id, cahier_id, payload)

        history_changed = any(
            [
                before["statut_test"] != updated.statut_test,
                before["type_test"] != updated.type_test,
                (before["commentaire"] or "") != (updated.commentaire or ""),
                (before["bug_titre_correction"] or "") != (updated.bug_titre_correction or ""),
                (before["bug_nom_tache"] or "") != (updated.bug_nom_tache or ""),
            ]
        )
        if history_changed:
            self.repo.add_cas_test_history(
                {
                    "cas_test_id": updated.id,
                    "cahier_id": cahier_id,
                    "changed_by_id": changed_by_id,
                    "old_statut_test": before["statut_test"],
                    "new_statut_test": updated.statut_test,
                    "old_type_test": before["type_test"],
                    "new_type_test": updated.type_test,
                    "old_commentaire": before["commentaire"],
                    "new_commentaire": updated.commentaire,
                    "old_bug_titre_correction": before["bug_titre_correction"],
                    "new_bug_titre_correction": updated.bug_titre_correction,
                    "old_bug_nom_tache": before["bug_nom_tache"],
                    "new_bug_nom_tache": updated.bug_nom_tache,
                }
            )

        if data.statut_test is not None:
            self.repo.recalculer_stats(cahier_id)
        return self._attach_user_story_display(updated)

    def list_cas_test_history(self, cahier_id: int, cas_id: int, projet_id: int) -> list:
        self._verifier_appartenance(cahier_id, projet_id)
        cas = self.repo.get_cas_test(cas_id, cahier_id)
        if not cas:
            raise HTTPException(status_code=404, detail="Cas de test introuvable.")
        return self.repo.list_cas_test_history(cas_id, cahier_id)

    def generer_suggestion_bug(
        self,
        cahier_id: int,
        cas_id: int,
        projet_id: int,
        user_id: Optional[int],
    ) -> dict:
        self._verifier_appartenance(cahier_id, projet_id)
        cas = self.repo.get_cas_test(cas_id, cahier_id)
        if not cas:
            raise HTTPException(status_code=404, detail="Cas de test introuvable.")

        titre, tache = self._generer_bug_fields_ia(cas, user_id)
        return {
            "bug_titre_correction": titre,
            "bug_nom_tache": tache,
        }

    def get_cahier(self, projet_id: int) -> CahierTestGlobal:
        cahier = self.repo.get_by_projet(projet_id)
        if not cahier:
            raise HTTPException(status_code=404, detail="Aucun cahier de tests généré pour ce projet.")
        return cahier

    def get_cahier_detail(self, projet_id: int) -> CahierTestGlobal:
        cahier = self.repo.get_detail_by_projet(projet_id)
        if not cahier:
            raise HTTPException(status_code=404, detail="Aucun cahier de tests généré pour ce projet.")
        self._attach_user_story_display_many(cahier.cas_tests or [])
        return cahier

    def get_statistiques(self, projet_id: int) -> dict:
        cahier = self.repo.get_by_projet(projet_id)
        if not cahier:
            raise HTTPException(status_code=404, detail="Aucun cahier de tests généré pour ce projet.")
        total       = cahier.nombre_total  or 0
        reussi      = cahier.nombre_reussi or 0
        echoue      = cahier.nombre_echoue or 0
        bloque      = cahier.nombre_bloque or 0
        non_execute = max(0, total - reussi - echoue - bloque)
        def pct(val: int) -> float:
            return round(val / total * 100, 1) if total > 0 else 0.0
        return {
            "version":            cahier.version,
            "nombre_total":       total,
            "nombre_reussi":      reussi,
            "nombre_echoue":      echoue,
            "nombre_bloque":      bloque,
            "nombre_non_execute": non_execute,
            "pct_reussi":         pct(reussi),
            "pct_echoue":         pct(echoue),
            "pct_bloque":         pct(bloque),
            "pct_non_execute":    pct(non_execute),
        }

    @staticmethod
    def _increment_report_minor_version(current_version: Optional[str]) -> str:
        major, minor, patch = 1, 0, 0
        if current_version:
            match = re.fullmatch(r"(\d+)\.(\d+)\.(\d+)", current_version.strip())
            if match:
                major, minor, patch = map(int, match.groups())
        return f"{major}.{minor + 1}.0"

    @staticmethod
    def _compute_rapport_stats(cahier: CahierTestGlobal) -> dict:
        total = cahier.nombre_total or 0
        reussi = cahier.nombre_reussi or 0
        echoue = cahier.nombre_echoue or 0
        bloque = cahier.nombre_bloque or 0
        executes = max(0, reussi + echoue + bloque)
        taux = round((reussi / executes) * 100, 2) if executes > 0 else 0.0
        return {
            "total": total,
            "reussi": reussi,
            "echoue": echoue,
            "bloque": bloque,
            "executes": executes,
            "taux_reussite": taux,
        }

    def _generer_rapport_qa_ia(self, cahier: CahierTestGlobal, user_id: Optional[int]) -> dict:
        stats = self._compute_rapport_stats(cahier)
        prompt = (
            "Genere un rapport QA base sur ces statistiques de cahier de tests global.\n"
            f"Version cahier: {cahier.version}\n"
            f"Total tests: {stats['total']}\n"
            f"Tests executes: {stats['executes']}\n"
            f"Tests reussis: {stats['reussi']}\n"
            f"Tests echoues: {stats['echoue']}\n"
            f"Tests bloques: {stats['bloque']}\n"
            f"Taux de reussite: {stats['taux_reussite']}\n"
        )
        try:
            api_key = self._get_api_key_for_request(user_id)
            raw = self._appeler_openrouter(prompt, api_key, RAPPORT_QA_SYSTEM_PROMPT)
            match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, re.DOTALL)
            if match:
                payload = json.loads(match.group(1))
            else:
                start = raw.find("{")
                end = raw.rfind("}") + 1
                payload = json.loads(raw[start:end])
            return {
                "statut": str(payload.get("statut") or "brouillon"),
                "recommandations": str(payload.get("recommandations") or "").strip(),
                "tendance": str(payload.get("tendance") or "stable"),
                "indice_qualite": float(payload.get("indice_qualite") or 0.0),
                "nombre_anomalies_critiques": int(payload.get("nombre_anomalies_critiques") or 0),
            }
        except Exception:
            return {
                "statut": "brouillon",
                "recommandations": "Prioriser la correction des tests echoues et stabiliser les tests bloques.",
                "tendance": "stable",
                "indice_qualite": 0.0,
                "nombre_anomalies_critiques": 0,
            }

    def generer_rapport_qa(
        self,
        cahier_id: int,
        projet_id: int,
        user_id: int,
        mode_generation: str = "manuelle",
        version: Optional[str] = None,
        recommandations: Optional[str] = None,
    ) -> RapportQA:
        cahier = self._verifier_appartenance(cahier_id, projet_id)
        stats = self._compute_rapport_stats(cahier)

        ai_payload = None
        if mode_generation == "ai":
            ai_payload = self._generer_rapport_qa_ia(cahier, user_id)

        rapport = self.db.query(RapportQA).filter(RapportQA.cahierId == cahier_id).first()
        next_version = version or self._increment_report_minor_version(rapport.version if rapport else None)

        recommandations_finales = (
            ai_payload["recommandations"]
            if ai_payload and ai_payload.get("recommandations")
            else (recommandations or "")
        )
        statut_final = (ai_payload.get("statut") if ai_payload else "brouillon") or "brouillon"

        if not rapport:
            rapport = RapportQA(
                cahierId=cahier_id,
                version=next_version,
                dateGeneration=datetime.utcnow(),
                statut=statut_final,
                tauxReussite=stats["taux_reussite"],
                nombreTestsExecutes=stats["executes"],
                nombreTestsReussis=stats["reussi"],
                nombreTestsEchoues=stats["echoue"],
                recommandations=recommandations_finales,
            )
            self.db.add(rapport)
            self.db.flush()
        else:
            rapport.version = next_version
            rapport.dateGeneration = datetime.utcnow()
            rapport.statut = statut_final
            rapport.tauxReussite = stats["taux_reussite"]
            rapport.nombreTestsExecutes = stats["executes"]
            rapport.nombreTestsReussis = stats["reussi"]
            rapport.nombreTestsEchoues = stats["echoue"]
            rapport.recommandations = recommandations_finales

            if rapport.indicateurs is None:
                self.db.add(IndicateurQualite(rapportId=rapport.id))

        indicateur = rapport.indicateurs
        if indicateur is None:
            indicateur = IndicateurQualite(rapportId=rapport.id)
            self.db.add(indicateur)

        indice_qualite = (
            ai_payload["indice_qualite"]
            if ai_payload and ai_payload.get("indice_qualite")
            else round(stats["taux_reussite"] / 20, 2)
        )
        tendance = ai_payload["tendance"] if ai_payload else "stable"
        anomalies_critiques = ai_payload["nombre_anomalies_critiques"] if ai_payload else 0

        indicateur.tauxCouverture = None
        indicateur.tauxReussite = stats["taux_reussite"]
        indicateur.nombreAnomalies = max(0, stats["echoue"] + stats["bloque"])
        indicateur.nombreAnomaliesCritiques = max(0, anomalies_critiques)
        indicateur.indiceQualite = max(0.0, indice_qualite)
        indicateur.tendance = tendance

        self.db.commit()
        self.db.refresh(rapport)

        self.notification_service.notify_user(
            user_id=user_id,
            titre="Rapport QA genere",
            message=f"Le rapport QA v{rapport.version} a ete genere pour le cahier {cahier_id}.",
            notification_type=TypeNotification.REPORT_GENERATED,
            priorite="moyenne",
        )

        return rapport

    def get_rapport_qa(self, cahier_id: int, projet_id: int) -> RapportQA:
        self._verifier_appartenance(cahier_id, projet_id)
        rapport = self.db.query(RapportQA).filter(RapportQA.cahierId == cahier_id).first()
        if not rapport:
            raise HTTPException(status_code=404, detail="Aucun rapport QA généré pour ce cahier.")
        return rapport

    def update_rapport_qa(
        self,
        cahier_id: int,
        projet_id: int,
        user_id: int,
        payload: dict,
    ) -> RapportQA:
        rapport = self.get_rapport_qa(cahier_id, projet_id)

        if payload.get("statut") is not None:
            rapport.statut = payload["statut"]
        if payload.get("recommandations") is not None:
            rapport.recommandations = payload["recommandations"]

        rapport.version = payload.get("version") or self._increment_report_minor_version(rapport.version)
        rapport.dateGeneration = datetime.utcnow()

        self.db.commit()
        self.db.refresh(rapport)

        self.notification_service.notify_user(
            user_id=user_id,
            titre="Rapport QA modifie",
            message=f"Le rapport QA v{rapport.version} a ete mis a jour.",
            notification_type=TypeNotification.REPORT_GENERATED,
            priorite="moyenne",
        )

        return rapport

    def exporter_rapport_qa_pdf(self, cahier_id: int, projet_id: int) -> bytes:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="La bibliothèque reportlab est requise pour l'export PDF. Installez-la avec : pip install reportlab",
            )

        rapport = self.get_rapport_qa(cahier_id, projet_id)
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(f"Rapport QA v{rapport.version}", styles["Title"]))
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"Projet ID: {projet_id} | Cahier ID: {cahier_id}", styles["Normal"]))
        story.append(Paragraph(f"Date: {rapport.dateGeneration}", styles["Normal"]))
        story.append(Paragraph(f"Statut: {rapport.statut}", styles["Normal"]))
        story.append(Paragraph(f"Taux de reussite: {rapport.tauxReussite}%", styles["Normal"]))
        story.append(Paragraph(f"Tests executes: {rapport.nombreTestsExecutes}", styles["Normal"]))
        story.append(Paragraph(f"Tests reussis: {rapport.nombreTestsReussis}", styles["Normal"]))
        story.append(Paragraph(f"Tests echoues: {rapport.nombreTestsEchoues}", styles["Normal"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Recommandations", styles["Heading2"]))
        story.append(Paragraph((rapport.recommandations or "Aucune recommandation."), styles["Normal"]))

        doc.build(story)
        buf.seek(0)
        return buf.read()

    def exporter_rapport_qa_word(self, cahier_id: int, projet_id: int) -> bytes:
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="La bibliothèque python-docx est requise pour l'export Word. Installez-la avec : pip install python-docx",
            )

        rapport = self.get_rapport_qa(cahier_id, projet_id)

        doc = Document()
        doc.add_heading(f"Rapport QA v{rapport.version}", 0)
        doc.add_paragraph(f"Projet ID: {projet_id}")
        doc.add_paragraph(f"Cahier ID: {cahier_id}")
        doc.add_paragraph(f"Date: {rapport.dateGeneration}")
        doc.add_paragraph(f"Statut: {rapport.statut}")
        doc.add_paragraph(f"Taux de reussite: {rapport.tauxReussite}%")
        doc.add_paragraph(f"Tests executes: {rapport.nombreTestsExecutes}")
        doc.add_paragraph(f"Tests reussis: {rapport.nombreTestsReussis}")
        doc.add_paragraph(f"Tests echoues: {rapport.nombreTestsEchoues}")
        doc.add_heading("Recommandations", level=1)
        doc.add_paragraph(rapport.recommandations or "Aucune recommandation.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def list_cas_tests(self, cahier_id: int, projet_id: int) -> list:
        self._verifier_appartenance(cahier_id, projet_id)
        cases = self.repo.list_cas_tests(cahier_id)
        return self._attach_user_story_display_many(cases)

    def list_user_stories_for_cahier(self, projet_id: int) -> list[dict]:
        user_stories = (
            self.db.query(UserStory)
            .join(Epic, UserStory.epic_id == Epic.id)
            .join(Module, Epic.module_id == Module.id)
            .options(joinedload(UserStory.sprint), joinedload(UserStory.epic).joinedload(Epic.module))
            .filter(Module.projet_id == projet_id)
            .order_by(UserStory.id.asc())
            .all()
        )

        return [
            {
                "id": us.id,
                "reference": us.reference,
                "titre": us.titre,
                "sprint_nom": us.sprint.nom if us.sprint else None,
                "module_nom": us.epic.module.nom if (us.epic and us.epic.module) else None,
            }
            for us in user_stories
        ]

    def get_assignable_members(self, cahier_id: int, projet_id: int) -> list[dict]:
        self._verifier_appartenance(cahier_id, projet_id)
        projet = self.db.query(Projet).filter(Projet.id == projet_id).first()
        if not projet:
            raise HTTPException(status_code=404, detail="Projet introuvable.")

        allowed_role_codes = {
            ROLE_TESTEUR_QA,
            ROLE_DEVELOPPEUR,
            "TESTER_QA",
            "TESTEUR",
            "DEVELOPER",
        }

        def normalize_role_code(value: Optional[str]) -> str:
            if not value:
                return ""
            raw = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
            return raw.strip().upper().replace(" ", "_")

        members: list[dict] = []
        for member in projet.membres or []:
            role_code = member.role.code if member.role else ""
            role_name = member.role.nom if member.role else ""
            normalized_role_code = normalize_role_code(role_code)
            normalized_role_name = normalize_role_code(role_name)

            # Safety deny rule: never expose Scrum Master as assignable member.
            if "SCRUM" in normalized_role_code or "SCRUM" in normalized_role_name:
                continue

            if not member.actif or normalized_role_code not in allowed_role_codes:
                continue
            members.append(
                {
                    "id": member.id,
                    "nom": member.nom,
                    "email": member.email,
                    "role_code": normalized_role_code,
                }
            )

        return sorted(members, key=lambda x: (x["nom"] or "").lower())

    def importer_excel(
        self,
        cahier_id: int,
        projet_id: int,
        file_content: bytes,
        current_user: Utilisateur,
    ) -> dict:
        """Importe des mises à jour de cas de tests depuis un fichier Excel."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="La bibliothèque openpyxl est requise pour l'import Excel.",
            )

        self._verifier_appartenance(cahier_id, projet_id)

        role_code = current_user.role.code if current_user and current_user.role else ""
        if role_code not in {ROLE_TESTEUR_QA, ROLE_DEVELOPPEUR}:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Import autorisé uniquement pour Testeur QA ou Développeur.",
            )

        def normalize_text(value: Optional[str]) -> str:
            raw = unicodedata.normalize("NFKD", value or "").encode("ascii", "ignore").decode("ascii")
            return raw.strip().lower()

        def normalize_name(value: Optional[str]) -> str:
            return normalize_text(value)

        def extract_assignee_tokens(value: Optional[str]) -> set[str]:
            text = (value or "").strip()
            if not text:
                return set()

            tokens = {normalize_name(text)}
            match = re.match(r"^(.*?)\s*\(([^)]+)\)\s*$", text)
            if match:
                name_part = normalize_name(match.group(1))
                email_part = normalize_name(match.group(2))
                if name_part:
                    tokens.add(name_part)
                if email_part:
                    tokens.add(email_part)
            return {t for t in tokens if t}

        def clean_cell(value):
            if value is None:
                return None
            if isinstance(value, str):
                stripped = value.strip()
                return stripped if stripped != "" else None
            return value

        def parse_status(value) -> Optional[str]:
            if value is None:
                return None
            mapping = {
                "non execute": "Non exécuté",
                "non executee": "Non exécuté",
                "non execute": "Non exécuté",
                "reussi": "Réussi",
                "echoue": "Échoué",
                "bloque": "Bloqué",
            }
            key = normalize_text(str(value))
            return mapping.get(key)

        def parse_type_test(value) -> Optional[str]:
            if value is None:
                return None
            key = normalize_text(str(value))
            if key in {"manuel", "manuelle"}:
                return "Manuel"
            if key in {"automatise", "automatisee", "automated"}:
                return "Automatisé"
            return None

        def parse_int(value) -> Optional[int]:
            if value is None:
                return None
            if isinstance(value, int):
                return max(0, value)
            if isinstance(value, float):
                return max(0, int(value))
            text = str(value).strip()
            if not text:
                return None
            return max(0, int(float(text)))

        header_aliases = {
            "test_ref": ["test ref", "ref", "test_ref"],
            "test_case": ["test case", "cas de test"],
            "test_purpose": ["test purpose", "objectif", "objectif du test"],
            "scenario_test": ["scenario test", "scenario", "scenario de test"],
            "resultat_attendu": ["resultat attendu"],
            "resultat_obtenu": ["resultat obtenu"],
            "execution_time_seconds": ["duree execution (s)", "execution time (s)", "duration / execution time (sec)"],
            "fail_logs": ["fail logs", "logs d'erreur", "logs d erreur"],
            "type_test": ["type", "type test"],
            "statut_test": ["statut test", "statut"],
            "commentaire": ["commentaire", "comment"],
            "bug_titre_correction": ["titre de correction", "bug titre correction"],
            "bug_nom_tache": ["nom de tache", "nom de tâche", "bug nom tache"],
        }

        workbook = load_workbook(io.BytesIO(file_content), data_only=True)
        sheet = workbook["Cas de Tests"] if "Cas de Tests" in workbook.sheetnames else workbook.active

        raw_headers = [sheet.cell(row=1, column=idx).value for idx in range(1, sheet.max_column + 1)]
        normalized_headers = {
            idx: normalize_text(str(value)) if value is not None else ""
            for idx, value in enumerate(raw_headers, start=1)
        }

        col_index: dict[str, int] = {}
        for field_name, aliases in header_aliases.items():
            for idx, header in normalized_headers.items():
                if header in aliases:
                    col_index[field_name] = idx
                    break

        if "test_ref" not in col_index:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Colonne 'Test REF' introuvable dans le fichier Excel.",
            )

        imported_count = 0
        skipped_refs: list[str] = []
        errors: list[str] = []
        current_user_tokens = {
            normalize_name(current_user.nom),
            normalize_name(current_user.email),
        }

        for row in range(2, sheet.max_row + 1):
            test_ref_val = clean_cell(sheet.cell(row=row, column=col_index["test_ref"]).value)
            if not test_ref_val:
                continue

            test_ref = str(test_ref_val).strip()
            cas = (
                self.db.query(CasTest)
                .filter(CasTest.cahier_id == cahier_id, CasTest.test_ref == test_ref)
                .first()
            )

            if not cas:
                errors.append(f"Ligne {row} ({test_ref}): cas introuvable.")
                continue

            assignee_tokens = extract_assignee_tokens(cas.type_utilisateur)
            if not assignee_tokens.intersection({t for t in current_user_tokens if t}):
                skipped_refs.append(test_ref)
                continue

            payload: dict = {}

            for field_name in [
                "test_case",
                "test_purpose",
                "scenario_test",
                "resultat_attendu",
                "resultat_obtenu",
                "fail_logs",
                "commentaire",
                "bug_titre_correction",
                "bug_nom_tache",
            ]:
                idx = col_index.get(field_name)
                if not idx:
                    continue
                cell_value = clean_cell(sheet.cell(row=row, column=idx).value)
                if cell_value is not None:
                    payload[field_name] = str(cell_value)

            idx_time = col_index.get("execution_time_seconds")
            if idx_time:
                try:
                    parsed_time = parse_int(clean_cell(sheet.cell(row=row, column=idx_time).value))
                    if parsed_time is not None:
                        payload["execution_time_seconds"] = parsed_time
                except Exception:
                    errors.append(f"Ligne {row} ({test_ref}): durée d'exécution invalide.")
                    continue

            idx_type = col_index.get("type_test")
            if idx_type:
                parsed_type = parse_type_test(clean_cell(sheet.cell(row=row, column=idx_type).value))
                raw_type = clean_cell(sheet.cell(row=row, column=idx_type).value)
                if raw_type is not None and parsed_type is None:
                    errors.append(f"Ligne {row} ({test_ref}): type de test invalide.")
                    continue
                if parsed_type is not None:
                    payload["type_test"] = parsed_type

            idx_status = col_index.get("statut_test")
            if idx_status:
                parsed_status = parse_status(clean_cell(sheet.cell(row=row, column=idx_status).value))
                raw_status = clean_cell(sheet.cell(row=row, column=idx_status).value)
                if raw_status is not None and parsed_status is None:
                    errors.append(f"Ligne {row} ({test_ref}): statut invalide.")
                    continue
                if parsed_status is not None:
                    payload["statut_test"] = parsed_status

            if not payload:
                skipped_refs.append(test_ref)
                continue

            try:
                self.update_cas_test(
                    cahier_id=cahier_id,
                    cas_id=cas.id,
                    projet_id=projet_id,
                    data=UpdateCasTestRequest(**payload),
                    changed_by_id=current_user.id,
                )
                imported_count += 1
            except HTTPException as exc:
                errors.append(f"Ligne {row} ({test_ref}): {exc.detail}")
            except Exception as exc:
                errors.append(f"Ligne {row} ({test_ref}): {str(exc)}")

        return {
            "imported_count": imported_count,
            "skipped_count": len(skipped_refs),
            "error_count": len(errors),
            "skipped_refs": skipped_refs,
            "errors": errors,
        }

    # ─── Logique interne de génération ────────────────────────────────────

    def _run(self, generation_id: int) -> None:

        # ── Étape 1 : Démarrage ────────────────────────────────────────────
        self.ai_repo.update_status(generation_id, "processing", 5)
        self.ai_repo.add_log(generation_id, "init",
                             "Démarrage de la génération du cahier de tests…", 5)

        gen = self.ai_repo.get_by_id(generation_id)
        projet = self.db.query(Projet).filter(Projet.id == gen.projet_id).first()

        # ── Étape 2 : Récupérer sprints & user stories ─────────────────────
        self.ai_repo.add_log(generation_id, "reading_sprints",
                             "Récupération des sprints et user stories…", 15)
        self.ai_repo.update_progress(generation_id, 15)

        sprints_content = self._construire_contenu_sprints(gen.projet_id)

        nb_sprints = sprints_content.count("--- Sprint") + sprints_content.count("---")
        self.ai_repo.add_log(generation_id, "reading_sprints",
                             f"Données récupérées ({len(sprints_content)} caractères).", 25)
        self.ai_repo.update_progress(generation_id, 25)

        # ── Étape 3 : Envoi du prompt à l'IA ──────────────────────────────
        self.ai_repo.add_log(generation_id, "sending_prompt",
                             "Envoi du prompt au modèle IA…", 35)
        self.ai_repo.update_progress(generation_id, 35)

        raw_json = self._appeler_ia(
            projet.nom,
            projet.description or "",
            sprints_content,
            generation_id,
            gen.user_id,
        )

        self.ai_repo.update_progress(generation_id, 65)
        self.ai_repo.add_log(generation_id, "parsing",
                             "Réponse IA reçue — analyse du JSON…", 65)

        # ── Étape 4 : Validation JSON ──────────────────────────────────────
        parsed = self._parser_reponse(raw_json)
        nb_cas = len(parsed.cas_tests)

        self.ai_repo.update_progress(generation_id, 75)
        self.ai_repo.add_log(generation_id, "saving",
                             f"{nb_cas} cas de tests détectés — sauvegarde en base…", 75)

        # ── Étape 5 : Sauvegarde ───────────────────────────────────────────
        cahier = self.repo.get_by_projet(gen.projet_id)
        if not cahier:
            raise ValueError("Cahier de tests introuvable lors de la sauvegarde.")

        for i, cas in enumerate(parsed.cas_tests, start=1):
            sprint = (cas.sprint or "")[:100]
            module = (cas.module or "")[:200]
            sous_module = (cas.sous_module or "")[:200]
            test_ref = (cas.test_ref or f"TC-{i:03d}")[:30]
            test_case = (cas.test_case or "")[:500]
            test_purpose = cas.test_purpose or ""
            type_utilisateur = (cas.type_utilisateur or "")[:100]
            scenario_test = cas.scenario_test or ""
            resultat_attendu = cas.resultat_attendu or ""
            execution_time_seconds = self._parse_execution_time_seconds(cas.execution_time_seconds)
            type_test = cas.type_test if cas.type_test in ("Manuel", "Automatisé") else "Manuel"
            resolved_user_story_id = self._resolve_user_story_id(
                projet_id=gen.projet_id,
                user_story_id=cas.user_story_id,
                module=module,
                sous_module=sous_module,
                test_case=test_case,
            )
            self.repo.add_cas_test(
                cahier_id=cahier.id,
                user_story_id=resolved_user_story_id,
                sprint=sprint,
                module=module,
                sous_module=sous_module,
                test_ref=test_ref,
                test_case=test_case,
                test_purpose=test_purpose,
                type_utilisateur=type_utilisateur,
                scenario_test=scenario_test,
                resultat_attendu=resultat_attendu,
                execution_time_seconds=execution_time_seconds,
                type_test=type_test,
                ordre=i,
            )

        self.repo.recalculer_stats(cahier.id)

        # Passer le cahier en brouillon (génération terminée)
        cahier.statut = "brouillon"
        self.db.commit()

        # ── Étape 6 : Terminé ──────────────────────────────────────────────
        self.ai_repo.update_status(generation_id, "completed", 100)
        self.ai_repo.add_log(
            generation_id, "done",
            f"Génération terminée : {nb_cas} cas de tests créés.",
            100,
        )

    # ─── Export ───────────────────────────────────────────────────────────

    def exporter_excel(self, cahier_id: int, projet_id: int) -> bytes:
        """Génère un fichier Excel (.xlsx) du cahier de tests."""
        from openpyxl import Workbook
        from openpyxl.styles import (
            Alignment, Border, Font, PatternFill, Side
        )
        from openpyxl.utils import get_column_letter

        cahier = self.repo.get_detail(cahier_id)
        if not cahier or cahier.projet_id != projet_id:
            raise HTTPException(status_code=404, detail="Cahier introuvable.")

        wb = Workbook()

        # ── Feuille résumé ────────────────────────────────────────────────
        ws_resume = wb.active
        ws_resume.title = "Résumé"

        header_fill = PatternFill("solid", fgColor="1F4E79")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        label_font  = Font(bold=True, size=11)
        center      = Alignment(horizontal="center", vertical="center")

        ws_resume.merge_cells("A1:D1")
        ws_resume["A1"] = f"Cahier de Tests — {cahier.projet.nom if cahier.projet else ''}"
        ws_resume["A1"].font      = Font(bold=True, color="FFFFFF", size=14)
        ws_resume["A1"].fill      = header_fill
        ws_resume["A1"].alignment = center

        summary_data = [
            ("Version",                  cahier.version),
            ("Statut",                   cahier.statut.capitalize()),
            ("Date de génération",       cahier.date_generation.strftime("%d/%m/%Y %H:%M") if cahier.date_generation else ""),
            ("Nombre total de tests",    cahier.nombre_total),
            ("Nombre de tests réussis",  cahier.nombre_reussi),
            ("Nombre de tests échoués",  cahier.nombre_echoue),
            ("Nombre de tests bloqués",  cahier.nombre_bloque),
        ]
        for row_idx, (label, value) in enumerate(summary_data, start=2):
            ws_resume.cell(row=row_idx, column=1, value=label).font = label_font
            ws_resume.cell(row=row_idx, column=2, value=value)
        ws_resume.column_dimensions["A"].width = 30
        ws_resume.column_dimensions["B"].width = 30

        # ── Feuille cas de tests ──────────────────────────────────────────
        ws = wb.create_sheet("Cas de Tests")

        columns = [
            ("Sprint",             20),
            ("Module",             22),
            ("Sous-Module",        22),
            ("Test REF",           12),
            ("Test Case",          35),
            ("Test Purpose",       35),
            ("Type Utilisateur",   18),
            ("Scénario Test",      50),
            ("Résultat Attendu",   40),
            ("Résultat Obtenu",    40),
            ("Durée Exécution (s)", 18),
            ("Fail Logs",          30),
            ("Capture",            20),
            ("Date Création",      18),
            ("Type",               12),
            ("Statut Test",        16),
            ("Commentaire",        35),
        ]

        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # En-têtes
        for col_idx, (col_name, col_width) in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font      = header_font
            cell.fill      = header_fill
            cell.alignment = center
            cell.border    = border
            ws.column_dimensions[get_column_letter(col_idx)].width = col_width

        # Couleurs de statut
        status_fills = {
            "Réussi":      PatternFill("solid", fgColor="C6EFCE"),
            "Échoué":      PatternFill("solid", fgColor="FFC7CE"),
            "Bloqué":      PatternFill("solid", fgColor="FFEB9C"),
            "Non exécuté": PatternFill("solid", fgColor="DDDDDD"),
        }

        for row_idx, cas in enumerate(cahier.cas_tests, start=2):
            row_data = [
                cas.sprint or "",
                cas.module or "",
                cas.sous_module or "",
                cas.test_ref,
                cas.test_case,
                cas.test_purpose or "",
                cas.type_utilisateur or "",
                cas.scenario_test or "",
                cas.resultat_attendu or "",
                cas.resultat_obtenu or "",
                cas.execution_time_seconds if cas.execution_time_seconds is not None else "",
                cas.fail_logs or "",
                cas.capture or "",
                cas.date_creation.strftime("%d/%m/%Y") if cas.date_creation else "",
                cas.type_test,
                cas.statut_test,
                cas.commentaire or "",
            ]
            status_fill = status_fills.get(cas.statut_test, status_fills["Non exécuté"])
            for col_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border    = border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                # Colorier la colonne Statut Test
                if col_idx == 15:
                    cell.fill = status_fill

        ws.freeze_panes = "A2"

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()

    def exporter_word(self, cahier_id: int, projet_id: int) -> bytes:
        """Génère un fichier Word (.docx) du cahier de tests."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        cahier = self.repo.get_detail(cahier_id)
        if not cahier or cahier.projet_id != projet_id:
            raise HTTPException(status_code=404, detail="Cahier introuvable.")

        doc = Document()

        # ── Titre ─────────────────────────────────────────────────────────
        titre = doc.add_heading(f"Cahier de Tests — {cahier.projet.nom if cahier.projet else ''}", 0)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Version : {cahier.version}")
        doc.add_paragraph(f"Statut : {cahier.statut.capitalize()}")
        date_str = cahier.date_generation.strftime("%d/%m/%Y %H:%M") if cahier.date_generation else ""
        doc.add_paragraph(f"Date de génération : {date_str}")
        doc.add_paragraph("")

        # ── Résumé statistiques ───────────────────────────────────────────
        doc.add_heading("Résumé", level=1)
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = "Light List Accent 1"
        stats_data = [
            ("Nombre total de tests",   str(cahier.nombre_total)),
            ("Tests réussis",           str(cahier.nombre_reussi)),
            ("Tests échoués",           str(cahier.nombre_echoue)),
            ("Tests bloqués",           str(cahier.nombre_bloque)),
            ("Tests non exécutés",      str(cahier.nombre_total - cahier.nombre_reussi - cahier.nombre_echoue - cahier.nombre_bloque)),
        ]
        for i, (label, val) in enumerate(stats_data):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = val

        doc.add_paragraph("")
        doc.add_heading("Cas de Tests", level=1)

        # ── Tableau des cas de tests ──────────────────────────────────────
        headers = [
            "Sprint", "Module", "Sous-Module", "Test REF", "Test Case",
            "Test Purpose", "Type Utilisateur", "Scénario Test",
            "Résultat Attendu", "Résultat Obtenu", "Durée Exécution (s)", "Fail Logs",
            "Date Création", "Type", "Statut Test", "Commentaire",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"

        # En-tête
        hdr_row = table.rows[0]
        for i, col_name in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = col_name
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)

        # Données
        for cas in cahier.cas_tests:
            row = table.add_row()
            values = [
                cas.sprint or "",
                cas.module or "",
                cas.sous_module or "",
                cas.test_ref,
                cas.test_case,
                cas.test_purpose or "",
                cas.type_utilisateur or "",
                cas.scenario_test or "",
                cas.resultat_attendu or "",
                cas.resultat_obtenu or "",
                str(cas.execution_time_seconds) if cas.execution_time_seconds is not None else "",
                cas.fail_logs or "",
                cas.date_creation.strftime("%d/%m/%Y") if cas.date_creation else "",
                cas.type_test,
                cas.statut_test,
                cas.commentaire or "",
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
                run.font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def exporter_pdf(self, cahier_id: int, projet_id: int) -> bytes:
        """Génère un fichier PDF du cahier de tests via reportlab."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
            )
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="La bibliothèque reportlab est requise pour l'export PDF. "
                       "Installez-la avec : pip install reportlab",
            )

        cahier = self.repo.get_detail(cahier_id)
        if not cahier or cahier.projet_id != projet_id:
            raise HTTPException(status_code=404, detail="Cahier introuvable.")

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=landscape(A4),
            leftMargin=1 * cm,
            rightMargin=1 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )

        styles  = getSampleStyleSheet()
        story   = []

        # Titre
        titre_style = ParagraphStyle(
            "titre",
            parent=styles["Title"],
            fontSize=16,
            spaceAfter=8,
        )
        projet_nom = cahier.projet.nom if cahier.projet else ""
        story.append(Paragraph(f"Cahier de Tests — {projet_nom}", titre_style))
        story.append(Spacer(1, 0.3 * cm))

        # Infos
        info_style = ParagraphStyle("info", parent=styles["Normal"], fontSize=9)
        date_str   = cahier.date_generation.strftime("%d/%m/%Y %H:%M") if cahier.date_generation else ""
        story.append(Paragraph(f"<b>Version :</b> {cahier.version}    "
                                f"<b>Statut :</b> {cahier.statut.capitalize()}    "
                                f"<b>Date :</b> {date_str}", info_style))
        story.append(Spacer(1, 0.3 * cm))

        # Stats
        nb_non_ex = cahier.nombre_total - cahier.nombre_reussi - cahier.nombre_echoue - cahier.nombre_bloque
        story.append(Paragraph(
            f"<b>Total :</b> {cahier.nombre_total}   "
            f"<b>Réussis :</b> {cahier.nombre_reussi}   "
            f"<b>Échoués :</b> {cahier.nombre_echoue}   "
            f"<b>Bloqués :</b> {cahier.nombre_bloque}   "
            f"<b>Non exécutés :</b> {nb_non_ex}",
            info_style,
        ))
        story.append(Spacer(1, 0.5 * cm))

        # Tableau des cas de tests
        small_style = ParagraphStyle("small", parent=styles["Normal"], fontSize=7, leading=9)
        header_style = ParagraphStyle("hdr", parent=styles["Normal"], fontSize=7, leading=9, textColor=colors.white)

        col_headers = [
            "Sprint", "Module", "Sous-Module", "REF", "Cas de Test",
            "Type User", "Scénario", "Résultat Attendu",
            "Résultat Obtenu", "Date", "Type", "Statut",
        ]
        col_widths = [2.5*cm, 2.8*cm, 2.8*cm, 1.6*cm, 4.5*cm,
                      2*cm, 6*cm, 4.5*cm, 3.5*cm, 1.8*cm, 1.6*cm, 2*cm]

        table_data = [[Paragraph(h, header_style) for h in col_headers]]
        for cas in cahier.cas_tests:
            table_data.append([
                Paragraph(cas.sprint or "", small_style),
                Paragraph(cas.module or "", small_style),
                Paragraph(cas.sous_module or "", small_style),
                Paragraph(cas.test_ref, small_style),
                Paragraph(cas.test_case, small_style),
                Paragraph(cas.type_utilisateur or "", small_style),
                Paragraph((cas.scenario_test or "").replace("\n", "<br/>"), small_style),
                Paragraph(cas.resultat_attendu or "", small_style),
                Paragraph(cas.resultat_obtenu or "", small_style),
                Paragraph(cas.date_creation.strftime("%d/%m/%Y") if cas.date_creation else "", small_style),
                Paragraph(cas.type_test, small_style),
                Paragraph(cas.statut_test, small_style),
            ])

        status_colors = {
            "Réussi":      colors.HexColor("#C6EFCE"),
            "Échoué":      colors.HexColor("#FFC7CE"),
            "Bloqué":      colors.HexColor("#FFEB9C"),
            "Non exécuté": colors.HexColor("#DDDDDD"),
        }

        ts = TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
            ("FONTSIZE",    (0, 0), (-1, -1), 7),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
            ("GRID",        (0, 0), (-1, -1), 0.3, colors.grey),
            ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ])

        for row_idx, cas in enumerate(cahier.cas_tests, start=1):
            fill = status_colors.get(cas.statut_test, status_colors["Non exécuté"])
            ts.add("BACKGROUND", (11, row_idx), (11, row_idx), fill)

        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(ts)
        story.append(t)

        doc.build(story)
        buf.seek(0)
        return buf.read()

    # ─── Helpers privés ───────────────────────────────────────────────────

    def _verifier_appartenance(self, cahier_id: int, projet_id: int) -> CahierTestGlobal:
        cahier = self.repo.get_by_id(cahier_id)
        if not cahier or cahier.projet_id != projet_id:
            raise HTTPException(status_code=404, detail="Cahier introuvable.")
        return cahier

    def _construire_contenu_sprints(self, projet_id: int) -> str:
        """Construit un résumé textuel des sprints et user stories pour le prompt IA."""
        sprints = (
            self.db.query(Sprint)
            .options(
                joinedload(Sprint.userstories)
                .joinedload(UserStory.epic)
                .joinedload(Epic.module)
            )
            .filter(Sprint.projet_id == projet_id)
            .order_by(Sprint.dateDebut.asc())
            .all()
        )

        if not sprints:
            raise ValueError(
                "Aucun sprint trouvé pour ce projet. "
                "Veuillez d'abord créer des sprints et des user stories.",
            )

        lines: List[str] = []
        for sprint in sprints:
            lines.append(f"\n--- {sprint.nom} ---")
            if sprint.objectifSprint:
                lines.append(f"Objectif : {sprint.objectifSprint}")
            if not sprint.userstories:
                lines.append("  (aucune user story assignée)")
                continue
            for us in sprint.userstories:
                module_nom     = us.epic.module.nom if (us.epic and us.epic.module) else "N/A"
                epic_nom       = us.epic.titre if us.epic else "N/A"
                us_ref = us.reference or f"US-{us.id}"
                lines.append(f"  [US-{us.id}] Reference: {us_ref} | Module: {module_nom} | Epic: {epic_nom}")
                lines.append(f"    Titre       : {us.titre}")
                if us.description:
                    lines.append(f"    Description : {us.description}")
                if us.criteresAcceptation:
                    lines.append(f"    Critères d'acceptation : {us.criteresAcceptation}")

        return "\n".join(lines)

    def _appeler_ia(self, projet_nom: str, projet_description: str,
                    sprints_content: str, generation_id: int, user_id: Optional[int]) -> str:
        """Envoie le prompt à l'IA et retourne la réponse brute."""
        api_key = self._get_api_key_for_request(user_id)

        # Limiter la taille du contenu
        max_chars = 40_000
        if len(sprints_content) > max_chars:
            sprints_content = sprints_content[:max_chars] + "\n[... contenu tronqué ...]"

        full_prompt = USER_PROMPT_TEMPLATE.format(
            projet_nom=projet_nom,
            projet_description=projet_description,
            sprints_content=sprints_content,
        )

        max_retries = 3
        for attempt in range(max_retries):
            try:
                return self._appeler_openrouter(full_prompt, api_key)
            except Exception as exc:
                err_str = str(exc)
                is_quota = "429" in err_str or "quota" in err_str.lower() or "RESOURCE_EXHAUSTED" in err_str
                if not is_quota or attempt == max_retries - 1:
                    raise
                delay_match = re.search(r"retry_delay\s*\{\s*seconds:\s*(\d+)", err_str)
                wait = int(delay_match.group(1)) + 5 if delay_match else 30 * (2 ** attempt)
                self.ai_repo.add_log(
                    generation_id, "retrying",
                    f"Quota dépassé (429) — nouvelle tentative dans {wait}s "
                    f"(essai {attempt + 1}/{max_retries - 1})…",
                    35,
                )
                time.sleep(wait)

        raise RuntimeError("Échec après toutes les tentatives IA.")

    @staticmethod
    def _appeler_openrouter(full_prompt: str, api_key: str, system_prompt: str = SYSTEM_PROMPT) -> str:
        """Appel unique vers l'API OpenRouter (compatible OpenAI)."""
        import requests

        clean_api_key = api_key.strip().strip('"').strip("'")

        payload = {
            "model": AI_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": full_prompt},
            ],
            "temperature": 0.2,
            "max_tokens":  8192,
        }
        headers = {
            "Authorization": f"Bearer {clean_api_key}",
            "Content-Type":  "application/json",
        }
        resp = requests.post(AI_API_URL, json=payload, headers=headers, timeout=120)

        if resp.status_code == 429:
            raise Exception(f"429 Quota dépassé : {resp.text}")

        if resp.status_code == 401:
            raise ValueError(
                "Erreur OpenRouter 401 : clé API invalide/expirée ou non autorisée pour ce compte. "
                "Vérifiez la clé API active (personnelle ou plateforme)."
            )

        if not resp.ok:
            raise ValueError(f"Erreur OpenRouter {resp.status_code} : {resp.text}")

        data = resp.json()
        return data["choices"][0]["message"]["content"]

    @staticmethod
    def _fallback_bug_fields(cas: CasTest) -> tuple[str, str]:
        base_title = (cas.test_case or "Cas de test").strip()
        base_task = (cas.module or cas.sous_module or cas.test_ref or "Bug").strip()
        titre = f"Corriger: {base_title}"[:120]
        tache = f"Fix {base_task} - {cas.test_ref or 'TC'}"[:120]
        return titre, tache

    @staticmethod
    def _parser_bug_suggestion(raw: str) -> tuple[str, str]:
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, re.DOTALL)
        if match:
            json_str = match.group(1)
        else:
            start = raw.find("{")
            end = raw.rfind("}") + 1
            if start == -1 or end == 0:
                raise ValueError("Réponse IA invalide pour la suggestion de bug.")
            json_str = raw[start:end]

        data = json.loads(json_str)
        titre = str(data.get("bug_titre_correction") or "").strip()
        tache = str(data.get("bug_nom_tache") or "").strip()
        if not titre or not tache:
            raise ValueError("Champs manquants dans la suggestion de bug.")
        return titre[:120], tache[:120]

    def _generer_bug_fields_ia(self, cas: CasTest, user_id: Optional[int]) -> tuple[str, str]:
        try:
            api_key = self._get_api_key_for_request(user_id)
            prompt = (
                "Propose un titre de correction et un nom de tâche pour ce cas de test.\n"
                f"Test REF: {cas.test_ref or ''}\n"
                f"Test case: {cas.test_case or ''}\n"
                f"Module: {cas.module or ''}\n"
                f"Sous-module: {cas.sous_module or ''}\n"
                f"Résultat obtenu: {cas.resultat_obtenu or ''}\n"
                f"Fail logs: {cas.fail_logs or ''}\n"
                f"Commentaire: {cas.commentaire or ''}\n"
            )
            raw = self._appeler_openrouter(prompt, api_key, BUG_SUGGESTION_SYSTEM_PROMPT)
            return self._parser_bug_suggestion(raw)
        except Exception:
            return self._fallback_bug_fields(cas)

    @staticmethod
    def _parser_reponse(raw: str) -> AICahierResponse:
        """Extrait et valide le JSON retourné par l'IA."""
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, re.DOTALL)
        if match:
            json_str = match.group(1)
        else:
            start = raw.find("{")
            end   = raw.rfind("}") + 1
            if start == -1 or end == 0:
                raise ValueError(
                    "L'IA n'a pas retourné de JSON valide. Relancez la génération."
                )
            json_str = raw[start:end]

        try:
            data = json.loads(json_str)
        except json.JSONDecodeError as e:
            raise ValueError(f"JSON malformé retourné par l'IA : {e}")

        return AICahierResponse(**data)
