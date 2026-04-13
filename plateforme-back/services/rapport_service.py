import io
import json
import re
from datetime import datetime
from typing import Optional

import requests
from fastapi import HTTPException
from sqlalchemy import func
from sqlalchemy.orm import Session

from core.config import AI_API_KEY, AI_API_URL, AI_MODEL
from core.rbac.constants import (
    ROLE_DEVELOPPEUR,
    ROLE_PRODUCT_OWNER,
    ROLE_SCRUM_MASTER,
    ROLE_TESTEUR_QA,
)
from models.cahier_test_global import CahierTestGlobal, CasTest, CasTestHistory
from models.notification import TypeNotification
from models.rapports import IndicateurQualite, RapportQA, RecommandationQualite
from models.scrum import Projet
from services.api_key_service import APIKeyService
from services.notification_service import NotificationService

RAPPORT_QA_SYSTEM_PROMPT = """\
Tu es un lead QA expert.

Tu recois les metriques d'execution d'un cahier de tests.
Ta mission est de generer une analyse qualite claire, professionnelle et actionnable.

Regles importantes:
- Sois factuel et base-toi uniquement sur les donnees fournies
- Evite toute duplication dans les recommandations
- Regroupe les idees similaires en une seule recommandation
- Priorise les actions (haute > moyenne > basse)
- Donne des recommandations concretes, precises et orientees execution
- Redige de maniere professionnelle (niveau entreprise SaaS)
- Mets en avant les risques (stabilite, regression, couverture)
- Priorise les problemes critiques (auth, paiement, securite, permissions, token)

Retourne UNIQUEMENT un JSON valide (sans texte autour):
{
    \"statut\": \"brouillon|valide\",
    \"recommandations\": \"resume global de la qualite + actions principales\",
    \"indicateur_qualite\": {
        \"taux_couverture\": 0.0,
        \"taux_reussite\": 0.0,
        \"nombre_anomalies\": 0,
        \"nombre_anomalies_critiques\": 0,
        \"indice_qualite\": 0.0,
        \"tendance\": \"amelioration|stable|degradation\"
    },
    \"recommandations_qualite\": [
        {
            \"titre\": \"texte court\",
            \"description\": \"action detaillee et concrete\",
            \"categorie\": \"stabilite|fiabilite|couverture|performance|process\",
            \"priorite\": \"haute|moyenne|basse\",
            \"impact\": 0.0
        }
    ]
}

Contraintes:
- indicateur_qualite.indice_qualite est un float entre 0.0 et 10.0
- recommandations_qualite[].impact est un float entre 0.0 et 1.0
- nombre_anomalies et nombre_anomalies_critiques sont des entiers >= 0
- Fournis entre 3 et 6 recommandations_qualite, sans duplication
- Trie recommandations_qualite par priorite (haute -> moyenne -> basse), puis impact desc
- Si nombre_anomalies_critiques > 0, la tendance ne peut pas etre amelioration
- Si taux_reussite est faible ou si tests critiques echoues/bloques, les premieres recommandations doivent etre de priorite haute

Compatibilite:
- Si tu utilises un champ \"analyse\", son contenu doit etre identique a \"recommandations\".
"""


class RapportService:
    def __init__(self, db: Session):
        self.db = db
        self.notification_service = NotificationService(db)
        self.api_key_service = APIKeyService(db)

    def _verifier_appartenance(self, cahier_id: int, projet_id: int) -> CahierTestGlobal:
        cahier = (
            self.db.query(CahierTestGlobal)
            .filter(CahierTestGlobal.id == cahier_id, CahierTestGlobal.projet_id == projet_id)
            .first()
        )
        if not cahier:
            raise HTTPException(status_code=404, detail="Cahier introuvable.")
        return cahier

    @staticmethod
    def _increment_report_minor_version(current_version: Optional[str]) -> str:
        major, minor, patch = 1, 0, 0
        if current_version:
            match = re.fullmatch(r"(\d+)\.(\d+)\.(\d+)", current_version.strip())
            if match:
                major, minor, patch = map(int, match.groups())
        return f"{major}.{minor + 1}.0"

    @staticmethod
    def _is_critical_case(cas: CasTest) -> bool:
        critical_keywords = (
            "auth",
            "login",
            "paiement",
            "payment",
            "securite",
            "security",
            "permission",
            "token",
            "checkout",
        )
        haystack = " ".join(
            [
                cas.test_case or "",
                cas.test_purpose or "",
                cas.module or "",
                cas.sous_module or "",
            ]
        ).lower()
        return any(keyword in haystack for keyword in critical_keywords)

    def _compute_rapport_stats(self, cahier: CahierTestGlobal) -> dict:
        total = cahier.nombre_total or 0
        reussi = cahier.nombre_reussi or 0
        echoue = cahier.nombre_echoue or 0
        bloque = cahier.nombre_bloque or 0
        executes = max(0, reussi + echoue + bloque)
        non_executes = max(0, total - executes)

        taux_reussite = round((reussi / executes) * 100, 2) if executes > 0 else 0.0
        taux_couverture = round((executes / total) * 100, 2) if total > 0 else 0.0

        tests = list(cahier.cas_tests or [])
        critical_failed = sum(
            1
            for cas in tests
            if cas.statut_test in ("Echoue", "Échoué", "Bloque", "Bloqué") and self._is_critical_case(cas)
        )

        bug_recurrence_rows = (
            self.db.query(CasTestHistory.cas_test_id, func.count(CasTestHistory.id).label("events"))
            .filter(
                CasTestHistory.cahier_id == cahier.id,
                CasTestHistory.new_statut_test.in_(["Echoue", "Échoué", "Bloque", "Bloqué"]),
            )
            .group_by(CasTestHistory.cas_test_id)
            .all()
        )
        bug_recurrence = int(sum(max(0, int(row.events) - 1) for row in bug_recurrence_rows))

        anomalies_total = max(0, echoue + bloque)

        return {
            "total": total,
            "reussi": reussi,
            "echoue": echoue,
            "bloque": bloque,
            "executes": executes,
            "non_executes": non_executes,
            "taux_reussite": taux_reussite,
            "taux_couverture": taux_couverture,
            "critical_failed": critical_failed,
            "bug_recurrence": bug_recurrence,
            "anomalies_total": anomalies_total,
        }

    def _build_manual_recommendations(self, stats: dict) -> dict:
        tendance = "stable"
        if stats["taux_reussite"] < 60 or stats["critical_failed"] >= 3:
            tendance = "degradation"
        elif stats["taux_reussite"] >= 85 and stats["bug_recurrence"] == 0:
            tendance = "amelioration"

        recommandations = []
        recommandations_qualite = []

        if stats["echoue"] > 0:
            recommandations.append(
                f"Traiter prioritairement {stats['echoue']} test(s) en echec et verifier les regressions associees."
            )
            recommandations_qualite.append(
                {
                    "titre": "Corriger les echecs de tests",
                    "description": "Prioriser les tests en echec avec analyse de cause racine.",
                    "categorie": "fiabilite",
                    "priorite": "haute",
                    "impact": 0.9,
                }
            )

        if stats["critical_failed"] > 0:
            recommandations.append(
                f"Stabiliser les parcours critiques ({stats['critical_failed']} test(s) critiques en echec/bloques)."
            )
            recommandations_qualite.append(
                {
                    "titre": "Stabiliser les parcours critiques",
                    "description": "Executer un plan de correction cible sur authentification, securite et paiement.",
                    "categorie": "stabilite",
                    "priorite": "haute",
                    "impact": 0.95,
                }
            )

        if stats["bug_recurrence"] > 0:
            recommandations.append(
                f"Reduire la recurrence des bugs ({stats['bug_recurrence']} recurrence(s) detectee(s))."
            )
            recommandations_qualite.append(
                {
                    "titre": "Reduire la recurrence des bugs",
                    "description": "Ajouter des tests de non-regression et renforcer la revue de correctifs.",
                    "categorie": "process",
                    "priorite": "moyenne",
                    "impact": 0.7,
                }
            )

        if stats["taux_couverture"] < 80:
            recommandations.append(
                f"Augmenter la couverture d'execution (actuel: {stats['taux_couverture']}%)."
            )
            recommandations_qualite.append(
                {
                    "titre": "Augmenter la couverture de tests",
                    "description": "Executer les cas non lances et completer les scenarios manquants.",
                    "categorie": "couverture",
                    "priorite": "moyenne",
                    "impact": 0.65,
                }
            )

        if not recommandations:
            recommandations.append("Le niveau qualite est satisfaisant. Continuer la surveillance des tests critiques.")
            recommandations_qualite.append(
                {
                    "titre": "Maintenir la qualite actuelle",
                    "description": "Conserver la strategie de test actuelle et monitorer les evolutions.",
                    "categorie": "fiabilite",
                    "priorite": "basse",
                    "impact": 0.4,
                }
            )

        indice_qualite = max(0.0, min(10.0, round((stats["taux_reussite"] * 0.7 + stats["taux_couverture"] * 0.3) / 10, 2)))

        return {
            "statut": "brouillon",
            "recommandations": "\n".join(recommandations),
            "taux_couverture": stats["taux_couverture"],
            "taux_reussite": stats["taux_reussite"],
            "nombre_anomalies": stats["anomalies_total"],
            "tendance": tendance,
            "indice_qualite": indice_qualite,
            "nombre_anomalies_critiques": stats["critical_failed"],
            "recommandations_qualite": recommandations_qualite,
        }

    def _get_api_key_for_request(self, user_id: Optional[int]) -> str:
        if user_id:
            custom = self.api_key_service.get_api_key_for_user(user_id)
            if custom:
                return custom
        if not AI_API_KEY:
            raise HTTPException(status_code=500, detail="Aucune cle API configuree pour la generation IA.")
        return AI_API_KEY

    def _get_rapport_notification_recipients(self, projet_id: int) -> list[int]:
        projet = self.db.query(Projet).filter(Projet.id == projet_id).first()
        if not projet:
            return []

        allowed_roles = {
            ROLE_TESTEUR_QA,
            ROLE_DEVELOPPEUR,
            ROLE_SCRUM_MASTER,
            ROLE_PRODUCT_OWNER,
        }

        recipients: set[int] = set()
        for member in projet.membres or []:
            role_code = member.role.code if member and member.role else None
            if member and member.actif and role_code in allowed_roles:
                recipients.add(member.id)

        if projet.product_owner and projet.product_owner.actif:
            recipients.add(projet.product_owner.id)

        return list(recipients)

    def _notify_rapport_stakeholders(
        self,
        projet_id: int,
        actor_user_id: int,
        titre: str,
        message: str,
    ) -> None:
        recipients = self._get_rapport_notification_recipients(projet_id)
        if not recipients:
            # Fallback minimal si aucun membre de projet n'est resolu.
            self.notification_service.notify_user(
                user_id=actor_user_id,
                titre=titre,
                message=message,
                notification_type=TypeNotification.REPORT_GENERATED,
                priorite="moyenne",
            )
            return

        self.notification_service.notify_users(
            user_ids=recipients,
            titre=titre,
            message=message,
            notification_type=TypeNotification.REPORT_GENERATED,
            priorite="moyenne",
        )

    def _generer_rapport_qa_ia(self, cahier: CahierTestGlobal, stats: dict, user_id: Optional[int]) -> dict:
        prompt = (
            "Genere un rapport QA base sur les donnees d'execution du cahier de tests global.\n"
            f"Version cahier: {cahier.version}\n"
            f"Total tests: {stats['total']}\n"
            f"Tests executes: {stats['executes']}\n"
            f"Tests non executes: {stats['non_executes']}\n"
            f"Tests reussis: {stats['reussi']}\n"
            f"Tests echoues: {stats['echoue']}\n"
            f"Tests bloques: {stats['bloque']}\n"
            f"Taux de reussite: {stats['taux_reussite']}\n"
            f"Couverture des tests: {stats['taux_couverture']}\n"
            f"Tests critiques en echec/bloques: {stats['critical_failed']}\n"
            f"Recurrence de bugs: {stats['bug_recurrence']}\n"
        )

        try:
            api_key = self._get_api_key_for_request(user_id)
            payload = {
                "model": AI_MODEL,
                "messages": [
                    {"role": "system", "content": RAPPORT_QA_SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.2,
                "max_tokens": 1200,
            }
            clean_api_key = api_key.strip().strip('"').strip("'")
            headers = {
                "Authorization": f"Bearer {clean_api_key}",
                "Content-Type": "application/json",
            }
            resp = requests.post(AI_API_URL, json=payload, headers=headers, timeout=120)
            if not resp.ok:
                raise ValueError(f"Erreur IA {resp.status_code}: {resp.text}")
            content = resp.json()["choices"][0]["message"]["content"]
            match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", content, re.DOTALL)
            if match:
                parsed = json.loads(match.group(1))
            else:
                start = content.find("{")
                end = content.rfind("}") + 1
                parsed = json.loads(content[start:end])

            items = parsed.get("recommandations_qualite")
            if not isinstance(items, list):
                alt_items = parsed.get("recommandations")
                items = alt_items if isinstance(alt_items, list) else []
            normalized_items = []
            for item in items[:6]:
                normalized_items.append(
                    {
                        "titre": str(item.get("titre") or "Action qualite"),
                        "description": str(item.get("description") or ""),
                        "categorie": str(item.get("categorie") or "fiabilite"),
                        "priorite": str(item.get("priorite") or "moyenne"),
                        "impact": float(item.get("impact") or 0.5),
                    }
                )

            indicateur_payload = parsed.get("indicateur_qualite") or {}
            tendances_alias = {
                "improving": "amelioration",
                "improvement": "amelioration",
                "degrading": "degradation",
                "declining": "degradation",
            }
            tendance_value = str(
                indicateur_payload.get("tendance")
                or parsed.get("tendance")
                or "stable"
            ).strip().lower()
            tendance_value = tendances_alias.get(tendance_value, tendance_value)

            summary_text = parsed.get("recommandations")
            if not isinstance(summary_text, str):
                summary_text = parsed.get("analyse") if isinstance(parsed.get("analyse"), str) else ""

            return {
                "statut": str(parsed.get("statut") or "brouillon"),
                "recommandations": str(summary_text or "").strip(),
                "taux_couverture": float(indicateur_payload.get("taux_couverture") or stats["taux_couverture"]),
                "taux_reussite": float(indicateur_payload.get("taux_reussite") or stats["taux_reussite"]),
                "nombre_anomalies": int(indicateur_payload.get("nombre_anomalies") or stats["anomalies_total"]),
                "tendance": tendance_value,
                "indice_qualite": float(
                    indicateur_payload.get("indice_qualite")
                    or parsed.get("indice_qualite")
                    or 0.0
                ),
                "nombre_anomalies_critiques": int(
                    indicateur_payload.get("nombre_anomalies_critiques")
                    or parsed.get("nombre_anomalies_critiques")
                    or 0
                ),
                "recommandations_qualite": normalized_items,
            }
        except Exception:
            return self._build_manual_recommendations(stats)

    @staticmethod
    def _sync_recommandations_qualite(rapport: RapportQA, items: list[dict]) -> None:
        rapport.recommandations_qualite.clear()
        for item in items:
            rapport.recommandations_qualite.append(
                RecommandationQualite(
                    titre=item.get("titre"),
                    description=item.get("description"),
                    categorie=item.get("categorie"),
                    priorite=item.get("priorite"),
                    impact=float(item.get("impact") or 0.0),
                    statut="ouverte",
                )
            )

    def generer_rapport_qa(
        self,
        cahier_id: int,
        projet_id: int,
        user_id: int,
        mode_generation: str = "manuelle",
        version: Optional[str] = None,
        recommandations: Optional[str] = None,
    ) -> RapportQA:
        cahier = self._verifier_appartenance(cahier_id, projet_id)
        stats = self._compute_rapport_stats(cahier)

        # Toujours tenter l'IA pour produire indicateur qualite + recommandations qualite.
        # En cas d'echec IA, _generer_rapport_qa_ia applique un fallback manuel.
        generated_payload = self._generer_rapport_qa_ia(cahier, stats, user_id)

        if recommandations and mode_generation == "manuelle":
            generated_payload["recommandations"] = recommandations

        rapport = self.db.query(RapportQA).filter(RapportQA.cahierId == cahier_id).first()
        next_version = version or self._increment_report_minor_version(rapport.version if rapport else None)

        if not rapport:
            rapport = RapportQA(
                cahierId=cahier_id,
                version=next_version,
                dateGeneration=datetime.utcnow(),
                statut=generated_payload.get("statut") or "brouillon",
                tauxReussite=stats["taux_reussite"],
                nombreTestsExecutes=stats["executes"],
                nombreTestsReussis=stats["reussi"],
                nombreTestsEchoues=stats["echoue"],
                recommandations=generated_payload.get("recommandations") or "",
            )
            self.db.add(rapport)
            self.db.flush()
        else:
            rapport.version = next_version
            rapport.dateGeneration = datetime.utcnow()
            rapport.statut = generated_payload.get("statut") or "brouillon"
            rapport.tauxReussite = stats["taux_reussite"]
            rapport.nombreTestsExecutes = stats["executes"]
            rapport.nombreTestsReussis = stats["reussi"]
            rapport.nombreTestsEchoues = stats["echoue"]
            rapport.recommandations = generated_payload.get("recommandations") or ""

        indicateur = rapport.indicateurs
        if indicateur is None:
            indicateur = IndicateurQualite(rapportId=rapport.id)
            self.db.add(indicateur)

        indicateur.tauxCouverture = float(generated_payload.get("taux_couverture") or stats["taux_couverture"])
        indicateur.tauxReussite = float(generated_payload.get("taux_reussite") or stats["taux_reussite"])
        indicateur.nombreAnomalies = int(generated_payload.get("nombre_anomalies") or stats["anomalies_total"])
        indicateur.nombreAnomaliesCritiques = max(0, int(generated_payload.get("nombre_anomalies_critiques") or 0))
        indicateur.indiceQualite = max(0.0, float(generated_payload.get("indice_qualite") or 0.0))
        indicateur.tendance = generated_payload.get("tendance") or "stable"

        self._sync_recommandations_qualite(
            rapport,
            generated_payload.get("recommandations_qualite") or [],
        )

        self.db.commit()
        self.db.refresh(rapport)

        self._notify_rapport_stakeholders(
            projet_id=projet_id,
            actor_user_id=user_id,
            titre="Rapport QA genere",
            message=f"Le rapport QA v{rapport.version} a ete genere pour le cahier {cahier_id}.",
        )

        return rapport

    def get_rapport_qa(self, cahier_id: int, projet_id: int) -> RapportQA:
        self._verifier_appartenance(cahier_id, projet_id)
        rapport = self.db.query(RapportQA).filter(RapportQA.cahierId == cahier_id).first()
        if not rapport:
            raise HTTPException(status_code=404, detail="Aucun rapport QA genere pour ce cahier.")
        return rapport

    def update_rapport_qa(self, cahier_id: int, projet_id: int, user_id: int, payload: dict) -> RapportQA:
        rapport = self.get_rapport_qa(cahier_id, projet_id)

        if payload.get("statut") is not None:
            rapport.statut = payload["statut"]
        if payload.get("recommandations") is not None:
            rapport.recommandations = payload["recommandations"]

        rapport.version = payload.get("version") or self._increment_report_minor_version(rapport.version)
        rapport.dateGeneration = datetime.utcnow()

        self.db.commit()
        self.db.refresh(rapport)

        self._notify_rapport_stakeholders(
            projet_id=projet_id,
            actor_user_id=user_id,
            titre="Rapport QA modifie",
            message=f"Le rapport QA v{rapport.version} a ete mis a jour.",
        )

        return rapport

    def exporter_rapport_qa_pdf(self, cahier_id: int, projet_id: int) -> bytes:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="La bibliotheque reportlab est requise pour l'export PDF. Installez-la avec : pip install reportlab",
            )

        rapport = self.get_rapport_qa(cahier_id, projet_id)

        def fmt_pct(value: Optional[float]) -> str:
            if value is None:
                return "0.0%"
            return f"{float(value):.1f}%"

        def fmt_num(value: Optional[float]) -> str:
            if value is None:
                return "0"
            return str(int(value))

        def fmt_date(value: Optional[datetime]) -> str:
            if not value:
                return "-"
            return value.strftime("%d/%m/%Y %H:%M")

        indic = rapport.indicateurs
        couverture = fmt_pct(indic.tauxCouverture if indic else 0.0)
        tendance = (indic.tendance if indic and indic.tendance else "stable").capitalize()
        indice_qualite = f"{float(indic.indiceQualite or 0.0):.1f}/10" if indic else "0.0/10"
        anomalies_critiques = fmt_num(indic.nombreAnomaliesCritiques if indic else 0)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=1.8 * cm,
            rightMargin=1.8 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.8 * cm,
        )
        styles = getSampleStyleSheet()
        story = []

        header = Table(
            [[f"Rapport QA v{rapport.version}", (rapport.statut or "brouillon").upper()]],
            colWidths=[doc.width * 0.75, doc.width * 0.25],
        )
        header.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#1F4E79")),
                ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#E2EFDA") if rapport.statut == "valide" else colors.HexColor("#FCE4D6")),
                ("TEXTCOLOR", (0, 0), (0, 0), colors.white),
                ("TEXTCOLOR", (1, 0), (1, 0), colors.HexColor("#375623") if rapport.statut == "valide" else colors.HexColor("#7F3F00")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 12),
                ("ALIGN", (1, 0), (1, 0), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ])
        )
        story.append(header)
        story.append(Spacer(1, 10))

        story.append(Paragraph(f"Projet ID: {projet_id} | Cahier ID: {cahier_id}", styles["Normal"]))
        story.append(Paragraph(f"Date de generation: {fmt_date(rapport.dateGeneration)}", styles["Normal"]))
        story.append(Spacer(1, 8))

        kpi = Table(
            [[
                f"Taux de reussite\n{fmt_pct(rapport.tauxReussite)}",
                f"Tests executes\n{fmt_num(rapport.nombreTestsExecutes)}",
                f"Tests reussis\n{fmt_num(rapport.nombreTestsReussis)}",
                f"Tests echoues\n{fmt_num(rapport.nombreTestsEchoues)}",
            ]],
            colWidths=[doc.width / 4] * 4,
        )
        kpi.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#D6E4F0")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1F4E79")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white),
                ("BOX", (0, 0), (-1, -1), 0, colors.white),
            ])
        )
        story.append(kpi)
        story.append(Spacer(1, 10))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2E75B6"), spaceAfter=6))
        story.append(Paragraph("Indicateurs qualite", styles["Heading2"]))
        indic_table = Table(
            [
                ["Couverture", couverture],
                ["Tendance", tendance],
                ["Indice qualite", indice_qualite],
                ["Anomalies critiques", anomalies_critiques],
            ],
            colWidths=[doc.width * 0.35, doc.width * 0.65],
        )
        indic_table.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#D6E4F0")),
                ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#F2F2F2")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1F4E79")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white),
                ("BOX", (0, 0), (-1, -1), 0, colors.white),
            ])
        )
        story.append(indic_table)

        story.append(Spacer(1, 12))
        story.append(Paragraph("Recommandations", styles["Heading2"]))
        story.append(Paragraph(rapport.recommandations or "Aucune recommandation.", styles["Normal"]))

        if rapport.recommandations_qualite:
            story.append(Spacer(1, 8))
            story.append(Paragraph("Actions recommandees", styles["Heading2"]))
            for rec in rapport.recommandations_qualite:
                line = f"- {rec.titre or 'Action qualite'} ({rec.priorite or 'moyenne'})"
                if rec.description:
                    line += f" : {rec.description}"
                story.append(Paragraph(line, styles["Normal"]))

        doc.build(story)
        buf.seek(0)
        return buf.read()

    def exporter_rapport_qa_word(self, cahier_id: int, projet_id: int) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="La bibliotheque python-docx est requise pour l'export Word. Installez-la avec : pip install python-docx",
            )

        rapport = self.get_rapport_qa(cahier_id, projet_id)
        indic = rapport.indicateurs

        def set_run_style(run, bold=False, size=11, color="000000"):
            run.bold = bold
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
            run.font.name = "Arial"

        def fmt_pct(value: Optional[float]) -> str:
            if value is None:
                return "0.0%"
            return f"{float(value):.1f}%"

        doc = Document()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(f"Rapport QA v{rapport.version}")
        set_run_style(run, bold=True, size=18, color="1F4E79")

        meta = doc.add_paragraph(f"Projet ID: {projet_id} | Cahier ID: {cahier_id}")
        meta2 = doc.add_paragraph(f"Date de generation: {rapport.dateGeneration}")
        for p in (meta, meta2):
            for r in p.runs:
                set_run_style(r, size=10, color="595959")

        kpi_table = doc.add_table(rows=2, cols=4)
        kpi_table.style = "Table Grid"
        headers = ["Taux de reussite", "Tests executes", "Tests reussis", "Tests echoues"]
        values = [
            fmt_pct(rapport.tauxReussite),
            str(rapport.nombreTestsExecutes or 0),
            str(rapport.nombreTestsReussis or 0),
            str(rapport.nombreTestsEchoues or 0),
        ]

        for idx in range(4):
            hcell = kpi_table.rows[0].cells[idx]
            vcell = kpi_table.rows[1].cells[idx]
            hcell.text = headers[idx]
            vcell.text = values[idx]
            for rr in hcell.paragraphs[0].runs:
                set_run_style(rr, bold=True, size=10, color="1F4E79")
            for rr in vcell.paragraphs[0].runs:
                set_run_style(rr, bold=True, size=12, color="000000")

        doc.add_paragraph()
        h_indic = doc.add_heading("Indicateurs qualite", level=1)
        for rr in h_indic.runs:
            set_run_style(rr, bold=True, size=14, color="1F4E79")

        indic_table = doc.add_table(rows=4, cols=2)
        indic_table.style = "Table Grid"
        indic_rows = [
            ("Couverture", fmt_pct(indic.tauxCouverture if indic else 0.0)),
            ("Tendance", (indic.tendance if indic and indic.tendance else "stable").capitalize()),
            ("Indice qualite", f"{float(indic.indiceQualite or 0.0):.1f}/10" if indic else "0.0/10"),
            ("Anomalies critiques", str(int(indic.nombreAnomaliesCritiques or 0)) if indic else "0"),
        ]
        for i, (label, value) in enumerate(indic_rows):
            indic_table.rows[i].cells[0].text = label
            indic_table.rows[i].cells[1].text = value
            for rr in indic_table.rows[i].cells[0].paragraphs[0].runs:
                set_run_style(rr, bold=True, size=10, color="1F4E79")
            for rr in indic_table.rows[i].cells[1].paragraphs[0].runs:
                set_run_style(rr, size=10, color="333333")

        h_recos = doc.add_heading("Recommandations", level=1)
        for rr in h_recos.runs:
            set_run_style(rr, bold=True, size=14, color="1F4E79")
        reco_para = doc.add_paragraph(rapport.recommandations or "Aucune recommandation.")
        for rr in reco_para.runs:
            set_run_style(rr, size=10, color="595959")

        if rapport.recommandations_qualite:
            h_actions = doc.add_heading("Actions recommandees", level=1)
            for rr in h_actions.runs:
                set_run_style(rr, bold=True, size=14, color="1F4E79")
            for rec in rapport.recommandations_qualite:
                line = f"- {rec.titre or 'Action qualite'} ({rec.priorite or 'moyenne'})"
                if rec.description:
                    line += f" : {rec.description}"
                p = doc.add_paragraph(line)
                for rr in p.runs:
                    set_run_style(rr, size=10, color="333333")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
